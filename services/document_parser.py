"""
文档解析服务 - 支持 PDF 和 Word 文档解析
"""
import os
from typing import Dict, List, Optional
from dataclasses import dataclass, field
import fitz  # PyMuPDF
from docx import Document

from config import SUPPORTED_EXTENSIONS, MAX_FILE_SIZE_MB, CHUNK_SIZE, CHUNK_OVERLAP


@dataclass
class ParsedDocument:
    """解析后的文档数据结构"""
    filename: str
    file_type: str
    title: str = ""
    content: str = ""
    chunks: List[str] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0


class DocumentParser:
    """文档解析器 - 支持 PDF 和 Word 格式"""
    
    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            ParsedDocument: 解析后的文档对象
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检查文件大小
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > MAX_FILE_SIZE_MB:
            raise ValueError(f"文件过大: {file_size_mb:.2f}MB，最大支持 {MAX_FILE_SIZE_MB}MB")
        
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {SUPPORTED_EXTENSIONS}")
        
        # 根据文件类型选择解析方法
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._parse_word(file_path)
        else:
            raise ValueError(f"未知的文件格式: {ext}")
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文档"""
        filename = os.path.basename(file_path)
        
        try:
            doc = fitz.open(file_path)
            
            # 提取文本内容
            full_text = ""
            for page_num, page in enumerate(doc):
                text = page.get_text()
                full_text += f"\n--- 第 {page_num + 1} 页 ---\n{text}"
            
            # 提取元数据
            metadata = doc.metadata or {}
            title = metadata.get("title", "") or self._extract_title_from_text(full_text)
            
            # 创建文本分块
            chunks = self._create_chunks(full_text)
            
            parsed_doc = ParsedDocument(
                filename=filename,
                file_type="pdf",
                title=title,
                content=full_text,
                chunks=chunks,
                metadata=metadata,
                page_count=len(doc),
                word_count=len(full_text.split())
            )
            
            doc.close()
            return parsed_doc
            
        except Exception as e:
            raise RuntimeError(f"解析 PDF 文件失败: {str(e)}")
    
    def _parse_word(self, file_path: str) -> ParsedDocument:
        """解析 Word 文档"""
        filename = os.path.basename(file_path)
        
        try:
            doc = Document(file_path)
            
            # 提取文本内容
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            full_text = "\n\n".join(paragraphs)
            
            # 提取标题（通常是第一个非空段落）
            title = self._extract_title_from_text(full_text)
            
            # 提取元数据
            metadata = {
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            }
            
            # 创建文本分块
            chunks = self._create_chunks(full_text)
            
            return ParsedDocument(
                filename=filename,
                file_type="word",
                title=title,
                content=full_text,
                chunks=chunks,
                metadata=metadata,
                page_count=len(doc.sections),
                word_count=len(full_text.split())
            )
            
        except Exception as e:
            raise RuntimeError(f"解析 Word 文件失败: {str(e)}")
    
    def _extract_title_from_text(self, text: str) -> str:
        """从文本中提取标题"""
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            # 跳过页码标记等
            if line and not line.startswith("---") and len(line) > 5:
                # 限制标题长度
                return line[:200] if len(line) > 200 else line
        return "未知标题"
    
    def _create_chunks(self, text: str) -> List[str]:
        """
        将文本分割成固定大小的块
        
        Args:
            text: 原始文本
            
        Returns:
            List[str]: 文本块列表
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            # 计算结束位置
            end = start + self.chunk_size
            
            if end >= text_length:
                # 最后一块
                chunks.append(text[start:].strip())
                break
            
            # 尝试在句子边界处分割
            # 向后查找句号、问号、感叹号等
            search_start = max(start + self.chunk_size - 100, start)
            best_split = end
            
            for sep in ["。", ".", "！", "!", "？", "?", "\n\n", "\n"]:
                pos = text.rfind(sep, search_start, end + 50)
                if pos != -1 and pos > search_start:
                    best_split = pos + len(sep)
                    break
            
            chunk = text[start:best_split].strip()
            if chunk:
                chunks.append(chunk)
            
            # 计算下一块的起始位置（带重叠）
            start = best_split - self.chunk_overlap
            if start < 0:
                start = 0
        
        return chunks
    
    def parse_from_bytes(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """
        从字节流解析文档（用于文件上传）
        
        Args:
            file_bytes: 文件字节流
            filename: 文件名
            
        Returns:
            ParsedDocument: 解析后的文档对象
        """
        import tempfile
        
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 先做一次大小校验，避免落盘/解析超大文件
        file_size_mb = len(file_bytes) / (1024 * 1024)
        if file_size_mb > MAX_FILE_SIZE_MB:
            raise ValueError(f"文件过大: {file_size_mb:.2f}MB，最大支持 {MAX_FILE_SIZE_MB}MB")
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        try:
            result = self.parse(tmp_path)
            return result
        finally:
            # 清理临时文件
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
